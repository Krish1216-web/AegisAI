import pytest
import os
import tempfile
import docx

from app.services.extractors.docx import DOCXExtractor

def test_docx_extractor_success():
    extractor = DOCXExtractor()

    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as f:
        temp_path = f.name
        
    try:
        # Create a real docx document
        doc = docx.Document()
        doc.add_paragraph("Some root level text.")
        doc.add_heading("Introduction", level=1)
        doc.add_paragraph("AegisAI document processing engine.")
        
        # Add table
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Quarter"
        table.cell(0, 1).text = "Revenue"
        table.cell(1, 0).text = "Q1"
        table.cell(1, 1).text = "$12000"
        
        doc.save(temp_path)
        
        # Extract
        res = extractor.extract(temp_path)
        
        assert "Introduction" in res.text
        assert "AegisAI document processing engine." in res.text
        assert "Quarter | Revenue" in res.text
        assert "Q1 | $12000" in res.text
        assert len(res.sections) == 2
        assert res.sections[0].title == "Root"
        assert res.sections[1].title == "Introduction"
        
    finally:
        if os.path.exists(temp_path):
            os.remove(temp_path)
